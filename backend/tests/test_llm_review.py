"""Tests for the Markdown extraction + ProfetAI review connector."""
from __future__ import annotations

import io
import json

import pytest
from docx import Document as DocxDocument

from app.models import DocumentStatus, GuardrailEvent, GuardrailType
from app.services import ingestion, llm as llm_module
from app.services.extraction import extract_text
from app.services.llm import (
    DocumentReview,
    LLMUnavailable,
    ProphetAILLM,
    _credentials_present,
    _extract_content,
    _parse_review,
    using_stub_llm,
)


def _docx_with_heading_and_table() -> bytes:
    doc = DocxDocument()
    doc.add_heading("專案目標", level=1)
    doc.add_paragraph("提升風控自動化。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Agent 名稱"
    table.cell(0, 1).text = "風控小幫手"
    table.cell(1, 0).text = "提案人 / 部門"
    table.cell(1, 1).text = "Alice / 風控部"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_extracts_markdown_heading_and_table():
    kind, md = extract_text("plan.docx", _docx_with_heading_and_table())
    assert kind == "docx"
    assert "# 專案目標" in md          # heading style → markdown heading
    assert "| Agent 名稱 | 風控小幫手 |" in md  # table → GFM table
    assert "| --- | --- |" in md


# ── response parsing ────────────────────────────────────────────────────


def test_extract_content_openai_shape():
    data = {"choices": [{"message": {"content": "hello"}}]}
    assert _extract_content(data) == "hello"


def test_extract_content_unknown_shape_raises():
    with pytest.raises(LLMUnavailable):
        _extract_content({"unexpected": True})


def test_parse_review_accepts_fenced_json():
    content = '```json\n{"verdict": "綠燈", "summary": "ok", "key_points": ["a"], "reasons": ["b"]}\n```'
    review = _parse_review(content, filename="x.docx")
    assert isinstance(review, DocumentReview)
    assert review.verdict == "綠燈"
    assert review.key_points == ["a"]


def test_parse_review_rejects_unknown_verdict():
    with pytest.raises(LLMUnavailable):
        _parse_review('{"verdict": "黃燈"}', filename="x.docx")


def test_parse_review_rejects_non_json():
    with pytest.raises(LLMUnavailable):
        _parse_review("the document looks fine to me", filename="x.docx")


# ── ProphetAILLM (monkeypatched transport) ─────────────────────────────────


def test_profetai_review_document_success(monkeypatch):
    client = ProphetAILLM()
    monkeypatch.setattr(
        client,
        "_chat",
        lambda **kw: '{"verdict": "紅燈", "summary": "缺章節", "reasons": ["缺風險"]}',
    )
    review = client.review_document(filename="p.docx", text="# 目標\n...")
    assert review.verdict == "紅燈"
    assert review.reasons == ["缺風險"]


def test_profetai_review_document_unavailable_propagates(monkeypatch):
    client = ProphetAILLM()

    def _boom(**kw):
        raise LLMUnavailable("connection refused")

    monkeypatch.setattr(client, "_chat", _boom)
    with pytest.raises(LLMUnavailable):
        client.review_document(filename="p.docx", text="x")


def test_chat_requires_host_credentials(monkeypatch):
    monkeypatch.delenv("COMPANY_LLM_API_KEY", raising=False)
    monkeypatch.delenv("COMPANY_LLM_AGENT", raising=False)
    with pytest.raises(LLMUnavailable):
        ProphetAILLM()._chat(prompt="hi")


def test_chat_builds_skill_compliant_request(monkeypatch):
    """model = agent id, content = block array, Bearer header (prophetai-api)."""
    monkeypatch.setenv("COMPANY_LLM_API_KEY", "ask_test123")
    monkeypatch.setenv("COMPANY_LLM_AGENT", "agent-xyz")

    captured = {}

    class _FakeResp:
        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        def read(self):
            return json.dumps(
                {"choices": [{"message": {"content": '{"verdict":"綠燈","summary":"ok"}'}}]}
            ).encode()

    class _FakeOpener:
        def open(self, req, timeout=None):
            captured["url"] = req.full_url
            captured["auth"] = req.headers.get("Authorization")
            captured["body"] = json.loads(req.data)
            return _FakeResp()

    monkeypatch.setattr(llm_module, "_no_proxy_ssl_off_opener", lambda: _FakeOpener())

    review = ProphetAILLM().review_document(filename="p.docx", text="# 目標")
    assert review.verdict == "綠燈"
    assert captured["auth"] == "Bearer ask_test123"
    assert captured["body"]["model"] == "agent-xyz"            # 規則 2
    content = captured["body"]["messages"][0]["content"]
    assert isinstance(content, list) and content[0]["type"] == "text"  # 規則 3


def test_using_stub_llm_toggles_with_credentials(monkeypatch):
    monkeypatch.delenv("COMPANY_LLM_API_KEY", raising=False)
    monkeypatch.delenv("COMPANY_LLM_AGENT", raising=False)
    assert using_stub_llm() is True
    monkeypatch.setenv("COMPANY_LLM_API_KEY", "ask_x")
    monkeypatch.setenv("COMPANY_LLM_AGENT", "agent-x")
    assert _credentials_present() is True
    assert using_stub_llm() is False


# ── ingestion: LLM failure → 無法審核 + Capability guardrail ───────────────


def test_ingest_llm_unavailable_marks_unreviewable(db, monkeypatch):
    class FailingLLM:
        def review_document(self, **kw):
            raise LLMUnavailable("api down")

    monkeypatch.setattr(ingestion, "llm", FailingLLM())
    doc = ingestion.ingest_file(
        db, filename="plan.txt", data="目標 範圍 時程 風險 資源 里程碑".encode("utf-8")
    )
    # File parsed fine, only the review failed.
    assert doc.status == DocumentStatus.DONE
    assert doc.llm_verdict == "無法審核"
    assert doc.created_project_id is None  # never auto-approve on failure
    events = db.query(GuardrailEvent).filter_by(guardrail_type=GuardrailType.CAPABILITY).all()
    assert len(events) == 1
